"""
Export full conversation transcript from PostgreSQL to Word document.
Reads all messages and creates a formatted table in .docx format.
"""
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import psycopg

DATABASE_URL = os.getenv("DATABASE_URL")
if not DATABASE_URL:
    print("ERROR: DATABASE_URL environment variable not set")
    sys.exit(1)

def get_all_messages():
    """Fetch all messages from database."""
    conn = psycopg.connect(DATABASE_URL)
    try:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT chat_id, role, content, timestamp
                FROM messages
                ORDER BY timestamp ASC
            """)
            return cur.fetchall()
    finally:
        conn.close()

def timestamp_to_datetime(ts):
    """Convert Unix timestamp to readable datetime."""
    return datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S")

def create_word_document(messages):
    """Create Word document with message table."""
    doc = Document()

    # Add title
    title = doc.add_heading("Robin Conversation Transcript", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    meta.add_run(f"\nTotal messages: {len(messages)}").italic = True

    doc.add_paragraph()

    # Create table: timestamp | chat_id | role | content
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Timestamp"
    hdr_cells[1].text = "Chat ID"
    hdr_cells[2].text = "Role"
    hdr_cells[3].text = "Message"

    # Format header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add message rows
    for chat_id, role, content, timestamp in messages:
        row_cells = table.add_row().cells
        row_cells[0].text = timestamp_to_datetime(timestamp)
        row_cells[1].text = chat_id
        row_cells[2].text = role

        # Handle potentially long content
        text_paragraph = row_cells[3].paragraphs[0]
        text_paragraph.text = content[:500]  # Truncate very long messages
        if len(content) > 500:
            text_paragraph.add_run(f"\n... (truncated, {len(content)} chars total)")

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)  # Timestamp
        row.cells[1].width = Inches(1.0)  # Chat ID
        row.cells[2].width = Inches(0.8)  # Role
        row.cells[3].width = Inches(2.5)  # Message

    return doc

def main():
    print("Fetching messages from database...")
    messages = get_all_messages()

    if not messages:
        print("No messages found in database.")
        return

    print(f"Found {len(messages)} messages")
    print("Creating Word document...")

    doc = create_word_document(messages)

    output_file = "robin_conversation.docx"
    doc.save(output_file)
    print(f"✓ Exported to {output_file}")
    print(f"  Date range: {timestamp_to_datetime(messages[0][3])} to {timestamp_to_datetime(messages[-1][3])}")

if __name__ == "__main__":
    main()
