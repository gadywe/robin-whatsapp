"""
File reading and creation tools for Robin.
Supports: txt, pdf, docx (read), audio (transcribe), docx+pdf (create)
"""
import io
import re
import tempfile
import os
import httpx
from typing import Optional


# ── READING ──────────────────────────────────────────────────────────────────

def read_text_file(content: bytes, encoding: str = "utf-8") -> str:
    try:
        return content.decode(encoding)
    except UnicodeDecodeError:
        return content.decode("windows-1255", errors="replace")


def read_pdf_file(content: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n\n".join(text_parts) if text_parts else "לא נמצא טקסט ב-PDF"


def read_docx_file(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs) if paragraphs else "המסמך ריק"


def extract_links(text: str) -> list[str]:
    pattern = r'https?://[^\s\)\]\>\"\'،,;]+'
    return re.findall(pattern, text)


def fetch_link_content(url: str, max_chars: int = 3000) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0 (compatible; Robin/1.0)"}
        resp = httpx.get(url, headers=headers, timeout=15, follow_redirects=True)
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "")
        if "text/html" in content_type:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(resp.text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            # Clean up whitespace
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            text = "\n".join(lines)
        else:
            text = resp.text
        return text[:max_chars] + ("..." if len(text) > max_chars else "")
    except Exception as e:
        return f"לא הצלחתי לפתוח את הלינק: {e}"


def process_file_by_mime(content: bytes, mime_type: str, filename: str = "") -> str:
    """Dispatch file processing based on MIME type."""
    mime = mime_type.lower()

    if "text/plain" in mime or filename.endswith(".txt"):
        text = read_text_file(content)
        links = extract_links(text)
        result = f"📄 תוכן הקובץ:\n{text[:4000]}"
        if links:
            result += f"\n\n🔗 לינקים שנמצאו: {len(links)}\n" + "\n".join(links[:5])
        return result

    elif "pdf" in mime or filename.endswith(".pdf"):
        text = read_pdf_file(content)
        links = extract_links(text)
        result = f"📑 תוכן ה-PDF:\n{text[:4000]}"
        if links:
            result += f"\n\n🔗 לינקים שנמצאו: {len(links)}\n" + "\n".join(links[:5])
        return result

    elif "word" in mime or "officedocument.wordprocessingml" in mime or filename.endswith(".docx"):
        text = read_docx_file(content)
        links = extract_links(text)
        result = f"📝 תוכן הקובץ:\n{text[:4000]}"
        if links:
            result += f"\n\n🔗 לינקים שנמצאו: {len(links)}\n" + "\n".join(links[:5])
        return result

    elif any(x in mime for x in ["audio", "mpeg", "mp4", "ogg", "webm"]) or \
         any(filename.endswith(ext) for ext in [".mp3", ".m4a", ".wav", ".ogg", ".aac"]):
        from transcribe import transcribe_audio_bytes
        text = transcribe_audio_bytes(content, mime_type=mime_type)
        return f"🎤 תמלול האודיו:\n{text}" if text else "לא הצלחתי לתמלל את הקובץ"

    else:
        return f"סוג קובץ לא נתמך: {mime_type} ({filename})"


# ── CREATING ─────────────────────────────────────────────────────────────────

def create_docx_bytes(content: str, title: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # RTL support
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if title:
        heading = doc.add_heading(title, 0)
        set_rtl(heading)

    for line in content.split("\n"):
        para = doc.add_paragraph(line)
        set_rtl(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_pdf_bytes(content: str, title: str = "") -> bytes:
    from fpdf import FPDF

    class HebrewPDF(FPDF):
        pass

    pdf = FPDF()
    pdf.add_page()

    # Use built-in font that supports basic Latin (Hebrew needs special font)
    # For proper Hebrew, we use a Unicode font
    font_path = os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf")

    if os.path.exists(font_path):
        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=12)
    else:
        pdf.set_font("Helvetica", size=12)

    if title:
        pdf.set_font_size(16)
        pdf.cell(0, 10, title, ln=True, align="R")
        pdf.set_font_size(12)
        pdf.ln(5)

    for line in content.split("\n"):
        try:
            pdf.multi_cell(0, 8, line, align="R")
        except Exception:
            pdf.multi_cell(0, 8, line.encode("latin-1", "replace").decode("latin-1"), align="R")

    return pdf.output()


# ── SENDING VIA WHATSAPP ──────────────────────────────────────────────────────

def upload_media_to_meta(file_bytes: bytes, mime_type: str, filename: str) -> str:
    """Upload a file to Meta and return the media_id."""
    from config import META_ACCESS_TOKEN, META_PHONE_NUMBER_ID
    url = f"https://graph.facebook.com/v19.0/{META_PHONE_NUMBER_ID}/media"
    headers = {"Authorization": f"Bearer {META_ACCESS_TOKEN}"}
    files = {
        "file": (filename, file_bytes, mime_type),
        "messaging_product": (None, "whatsapp"),
        "type": (None, mime_type),
    }
    resp = httpx.post(url, headers=headers, files=files, timeout=30)
    resp.raise_for_status()
    return resp.json()["id"]


def send_document_whatsapp(to: str, file_bytes: bytes, mime_type: str, filename: str, caption: str = ""):
    """Upload and send a document via WhatsApp."""
    from config import META_ACCESS_TOKEN, META_PHONE_NUMBER_ID
    media_id = upload_media_to_meta(file_bytes, mime_type, filename)
    url = f"https://graph.facebook.com/v19.0/{META_PHONE_NUMBER_ID}/messages"
    headers = {"Authorization": f"Bearer {META_ACCESS_TOKEN}", "Content-Type": "application/json"}
    payload = {
        "messaging_product": "whatsapp",
        "to": to,
        "type": "document",
        "document": {
            "id": media_id,
            "filename": filename,
            "caption": caption,
        }
    }
    resp = httpx.post(url, json=payload, headers=headers, timeout=30)
    resp.raise_for_status()
    return resp.json()
